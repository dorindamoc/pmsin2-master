from docx import Document
from docx.shared import Cm, Pt
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.text import WD_LINE_SPACING, WD_ALIGN_PARAGRAPH
import pandas as pd
import math
import glob
import gspread
from oauth2client.service_account import ServiceAccountCredentials
import logging
logging.basicConfig(level=logging.DEBUG) # DEBUG, INFO, WARN
import numpy as np
import math
from theone_con import from_theone
import sqlite3






class pm_doc:
    '''
    Document object with methods for creating and filling tables for a management plan
    Works in conjunction with a species object that has the requirs attributes
    Methods:
    - ftblA_sp - makes the table A for species
    - ftblB_sp - makes the table B for species
    - ftblA_cons - makes the table A for conservation
    - ftblB_cons - makes the table B for conservation
    - ftblC_cons - makes the table C for conservation
    - ftblD_cons - makes the table D for conservation
    - ftblM_1 - makes the table for matrix 1
    - ftblM_2 - makes the table for matrix 2
    - ftblM_3 - makes the table for matrix 3
    - ftblM_4 - makes the table for matrix 4
    - ftblM_5 - makes the table for matrix 5
    - ftblM_6 - makes the table for matrix 6
    - ftblM_7 - makes the table for matrix 7

    - "ftable_name"+_fill - fills that table. Require a species as argument. The species obj must have required attributes
    - fimg - inserts images tblA_sp
    - fmap - inserst maps in tblB_sp


    - sp_header - Add header with the species
    - empty_p - Add empty paraghraph

    - save(self, fn)
    '''
    def __init__(self):
        self.d = Document()
        styles = self.d.styles
        styles['Normal'].font.name = 'Times New Roman'
        styles['Normal'].font.size = Pt(11)


        poim_lat = styles.add_style('poim_lat', WD_STYLE_TYPE.CHARACTER)
        poim_lat.font.italic = True
        poim_lat.hidden = False
        poim_lat.quick_style = True
        poim_lat.priority = 1
        
        poim_h1 = styles.add_style('poim_h1', WD_STYLE_TYPE.PARAGRAPH)
        poim_h1.base_style = styles['Normal']
        poim_h1.font.bold = True
        poim_h1.font.size = Pt(18)
        poim_h1.hidden = False
        poim_h1.quick_style = True
        poim_h1.priority = 1
        poim_h1.paragraph_format.space_before = Pt(18)
        poim_h1.paragraph_format.space_after = Pt(6)
        
        poim_h2 = styles.add_style('poim_h2', WD_STYLE_TYPE.PARAGRAPH)
        poim_h2.base_style = styles['Normal']
        poim_h2.font.bold = True
        poim_h2.font.size = Pt(12)
        poim_h2.hidden = False
        poim_h2.quick_style = True
        poim_h2.priority = 1
        poim_h2.paragraph_format.space_before = Pt(12)
        poim_h2.paragraph_format.space_after = Pt(6)        

        poim_h3 = styles.add_style('poim_h3', WD_STYLE_TYPE.PARAGRAPH)
        poim_h3.base_style = styles['Normal']
        poim_h3.hidden = False
        poim_h3.quick_style = True
        poim_h3.priority = 1
        poim_h3.paragraph_format.space_before = Pt(12)
        poim_h3.paragraph_format.space_after = Pt(6)
        
        poim_tg = styles.add_style('poim_tg', WD_STYLE_TYPE.TABLE)
        poim_tg.base_style = styles['Table Grid']
        poim_tg.font.name = 'Times New Roman'
        poim_tg.hidden = False
        poim_tg.quick_style = True
        poim_tg.priority = 1
        poim_tg.paragraph_format.space_before = Pt(0)
        poim_tg.paragraph_format.space_after = Pt(0)
        
        poim_th = styles.add_style('poim_th', WD_STYLE_TYPE.PARAGRAPH)
        poim_th.base_style = styles['Normal']
        poim_th.font.bold = True
        poim_th.hidden = False
        poim_th.quick_style = True
        poim_th.priority = 1
        poim_th.paragraph_format.space_before = Pt(3)
        poim_th.paragraph_format.space_after = Pt(3)


    def empty_p(self):
        self.d.add_paragraph('')

    def chapter_h1(self, txt):
        ch1 = self.d.add_paragraph(txt)
        ch1.style = 'poim_h1'

    def chapter_h2(self, txt):
        ch1 = self.d.add_paragraph(txt)
        ch1.style = 'poim_h2'

    def chapter_h3(self, txt):
        ch1 = self.d.add_paragraph(txt)
        ch1.style = 'poim_h3'

    def sp_header(self, s):
        ph1 = self.d.add_paragraph(str(s.as1) + ' - ')
        ph1.style = 'poim_h1'
        ph1.add_run(s.as2 + ', ', 'poim_lat') 
        ph1.add_run(s.as3)


    #Tabel A for species description
    def ftblA_sp(self,chk,s):
        #Tabelul A
        tableA_title = self.d.add_paragraph('Tabelul A.	Date generale ale speciei')
        tableA_title.style = 'poim_h2'    
        self.tblA_sp = self.d.add_table(rows=8, cols=3)
        self.tblA_sp.rows[0].cells[0].text = 'Nr.'
        self.tblA_sp.rows[0].cells[0].style = 'poim_th'
        self.tblA_sp.rows[0].cells[1].text = 'Informație/Atribut'
        self.tblA_sp.rows[0].cells[2].text = 'Descriere'
        self.tblA_sp.rows[1].cells[0].text = '1'
        self.tblA_sp.rows[1].cells[1].text = 'Cod Specie - EUNIS'
        self.tblA_sp.rows[2].cells[0].text = '2'
        self.tblA_sp.rows[2].cells[1].text = 'Denumirea științifică'
        self.tblA_sp.rows[3].cells[0].text = '3'
        self.tblA_sp.rows[3].cells[1].text = 'Denumirea populară'
        self.tblA_sp.rows[4].cells[0].text = '4'
        self.tblA_sp.rows[4].cells[1].text = 'Descrierea speciei'
        self.tblA_sp.rows[5].cells[0].text = '5'
        self.tblA_sp.rows[5].cells[1].text = 'Perioade critice'
        self.tblA_sp.rows[6].cells[0].text = '6'
        self.tblA_sp.rows[6].cells[1].text = 'Cerințe de habitat'
        self.tblA_sp.rows[7].cells[0].text = '7'
        self.tblA_sp.rows[7].cells[1].text = 'Fotografii'

        self.tblA_sp.style = 'poim_tg'
        for cell in self.tblA_sp.columns[0].cells:
            cell.width = Cm(1)
        for cell in self.tblA_sp.columns[1].cells:
            cell.width = Cm(5)
        for cell in self.tblA_sp.columns[2].cells:
            cell.width = Cm(10)
        self.tblA_sp.autofit = False

        #Filler of Tabel A for species description
        if chk:
            self.tblA_sp.rows[1].cells[2].text = s.as1
            self.tblA_sp.rows[2].cells[2].paragraphs[0].add_run(s.as2, 'poim_lat')
            self.tblA_sp.rows[3].cells[2].text = s.as3
            #self.tblA_sp.rows[4].cells[2].text = s.as4
            self.tblA_sp.rows[4].cells[2].paragraphs[0].add_run('Descriere generală:')
            self.tblA_sp.rows[4].cells[2].paragraphs[0].style='poim_th'
            self.tblA_sp.rows[4].cells[2].add_paragraph(s.as4_gen)
            self.tblA_sp.rows[4].cells[2].add_paragraph('Descriere hrană:', 'poim_th')
            self.tblA_sp.rows[4].cells[2].add_paragraph(s.as4_hrana)
            self.tblA_sp.rows[4].cells[2].add_paragraph('Descriere cuibărit:', 'poim_th')
            self.tblA_sp.rows[4].cells[2].add_paragraph(s.as4_cuib)
            self.tblA_sp.rows[5].cells[2].text = s.as5
            self.tblA_sp.rows[6].cells[2].text = s.as6
            #self.tblA_sp.rows[7].cells[2].text = s.as7

            if s.as7 =='na':
                self.tblA_sp.rows[7].cells[2].text = 'Nu există imagine disponibilă cu această specie.'
            else:
                run_for_image = self.tblA_sp.rows[7].cells[2].paragraphs[0].add_run()
                run_for_image.add_picture(s.as7, width=Cm(9))
        
    #Tabel B for species description
    def ftblB_sp(self, chk, s):
        tableB_title = self.d.add_paragraph('Tabelul B.	Date specifice speciei la nivelul ariei naturale protejate')
        tableB_title.style = 'poim_h2' 
        self.tblB_sp = self.d.add_table(rows=11, cols=3)
        self.tblB_sp.rows[0].cells[0].text = 'Nr.'
        self.tblB_sp.rows[0].cells[1].text = 'Informație/Atribut'
        self.tblB_sp.rows[0].cells[2].text = 'Descriere'
        self.tblB_sp.rows[1].cells[0].text = '1'
        self.tblB_sp.rows[1].cells[1].text = 'Specia'
        self.tblB_sp.rows[2].cells[0].text = '2'
        self.tblB_sp.rows[2].cells[1].text = 'Informații specifice speciei'
        self.tblB_sp.rows[3].cells[0].text = '3'
        self.tblB_sp.rows[3].cells[1].text = 'Statutul de prezență [temporal]'
        self.tblB_sp.rows[4].cells[0].text = '4'
        self.tblB_sp.rows[4].cells[1].text = 'Statutul de prezență [spațial]'
        self.tblB_sp.rows[5].cells[0].text = '5'
        self.tblB_sp.rows[5].cells[1].text = 'Statutul de prezență [management]'
        self.tblB_sp.rows[6].cells[0].text = '6'
        self.tblB_sp.rows[6].cells[1].text = 'Abundență'
        self.tblB_sp.rows[7].cells[0].text = '7'
        self.tblB_sp.rows[7].cells[1].text = 'Perioada de colectare a datelor din teren'
        self.tblB_sp.rows[8].cells[0].text = '8'
        self.tblB_sp.rows[8].cells[1].text = 'Distribuția speciei [interpretare]'
        self.tblB_sp.rows[9].cells[0].text = '9'
        self.tblB_sp.rows[9].cells[1].text = 'Distribuția speciei [harta distribuției]'
        self.tblB_sp.rows[10].cells[0].text = '10'
        self.tblB_sp.rows[10].cells[1].text = 'Alte informații privind sursele de informații'

        self.tblB_sp.style = 'poim_tg'
        for cell in self.tblB_sp.columns[0].cells:
            cell.width = Cm(1)
        for cell in self.tblB_sp.columns[1].cells:
            cell.width = Cm(5)
        for cell in self.tblB_sp.columns[2].cells:
            cell.width = Cm(10)
        self.tblB_sp.autofit = False

        #Filler of Tabel B for species description
        if chk:
            self.tblB_sp.rows[1].cells[2].paragraphs[0].add_run(s.bs1, 'poim_lat')
            self.tblB_sp.rows[1].cells[2].paragraphs[0].add_run(s.bs11)
            self.tblB_sp.rows[2].cells[2].text = s.bs2
            self.tblB_sp.rows[3].cells[2].text = s.bs3
            self.tblB_sp.rows[4].cells[2].text = s.bs4
            self.tblB_sp.rows[5].cells[2].text = s.bs5
            self.tblB_sp.rows[6].cells[2].text = s.bs6
            self.tblB_sp.rows[7].cells[2].text = s.bs7
            self.tblB_sp.rows[8].cells[2].text = s.bs8
            #self.tblB_sp.rows[9].cells[2].text = s.bs9
            self.tblB_sp.rows[10].cells[2].text = s.bs10

            if s.bs9 =='na':
                self.tblB_sp.rows[9].cells[2].text = 'Nu există hartă disponibilă cu această specie.'
            else:
                run_for_map = self.tblB_sp.rows[9].cells[2].paragraphs[0].add_run()
                run_for_map.add_picture(s.bs9, width=Cm(9))

    #Tabel A for species conservation
    def ftblA_cons(self, chk, s):
        ftblA_title = self.d.add_paragraph('Tabelul A) Parametri pentru evaluarea stării de conservare a speciei din punct de vedere al populaţiei')
        ftblA_title.style = 'poim_h2' 
        self.tblA_cons = self.d.add_table(rows=19, cols=3)
            #Set table style and column dimensions
        self.tblA_cons.style = 'poim_tg'
        for cell in self.tblA_cons.columns[0].cells:
            cell.width = Cm(2)
        for cell in self.tblA_cons.columns[1].cells:
            cell.width = Cm(5)
        for cell in self.tblA_cons.columns[2].cells:
            cell.width = Cm(9)
        self.tblA_cons.autofit = False
            #Set table fixed content
        self.tblA_cons.rows[0].cells[0].text = 'Nr.'
        self.tblA_cons.rows[0].cells[1].text = 'Parametru'
        self.tblA_cons.rows[0].cells[2].text = 'Descriere'
        self.tblA_cons.rows[1].cells[0].text = 'A.1'
        self.tblA_cons.rows[1].cells[1].text = 'Specia'
        self.tblA_cons.rows[2].cells[0].text = 'A.2'
        self.tblA_cons.rows[2].cells[1].text = 'Statut de prezenţă temporală a speciilor'
        self.tblA_cons.rows[3].cells[0].text = 'A.3'
        self.tblA_cons.rows[3].cells[1].text = 'Mărimea populaţiei speciei în aria naturală protejată'
        self.tblA_cons.rows[4].cells[0].text = 'A.4'
        self.tblA_cons.rows[4].cells[1].text = 'Calitatea datelor referitoare la populaţia speciei din aria naturală protejată'
        self.tblA_cons.rows[5].cells[0].text = 'A.5'
        self.tblA_cons.rows[5].cells[1].text = 'Raportul dintre mărimea populaţiei speciei în aria naturală protejată şi mărimea populaţiei naţionale'
        self.tblA_cons.rows[6].cells[0].text = 'A.6'
        self.tblA_cons.rows[6].cells[1].text = 'Mărimea populaţiei speciei în aria naturală protejată comparata cu mărimea populaţiei naţionale'
        self.tblA_cons.rows[7].cells[0].text = 'A.7'
        self.tblA_cons.rows[7].cells[1].text = 'Mărimea reevaluată a populaţiei estimate în planul de management anterior'
        self.tblA_cons.rows[8].cells[0].text = 'A.8'
        self.tblA_cons.rows[8].cells[1].text = 'Mărimea populaţiei de referinţă pentru starea favorabilă în aria naturală protejată'
        self.tblA_cons.rows[9].cells[0].text = 'A.9'
        self.tblA_cons.rows[9].cells[1].text = 'Metodologia de apreciere a mărimii populaţiei de referinţă pentru starea favorabilă'
        self.tblA_cons.rows[10].cells[0].text = 'A.10'
        self.tblA_cons.rows[10].cells[1].text = 'Raportul dintre mărimea populaţiei de referinţă pentru starea favorabilă şi mărimea populaţiei actuale'
        self.tblA_cons.rows[11].cells[0].text = 'A.11'
        self.tblA_cons.rows[11].cells[1].text = 'Tendinţa actuală a mărimii populaţiei speciei'
        self.tblA_cons.rows[12].cells[0].text = 'A.12'
        self.tblA_cons.rows[12].cells[1].text = 'Calitatea datelor privind tendinţa actuală a mărimii populaţiei speciei'
        self.tblA_cons.rows[13].cells[0].text = 'A.13'
        self.tblA_cons.rows[13].cells[1].text = 'Magnitudinea tendinţei actuale a mărimii populaţiei speciei'
        self.tblA_cons.rows[14].cells[0].text = 'A.14'
        self.tblA_cons.rows[14].cells[1].text = 'Magnitudinea tendinţei actuale a mărimii populaţiei speciei exprimată prin calificative'
        self.tblA_cons.rows[15].cells[0].text = 'A.15'
        self.tblA_cons.rows[15].cells[1].text = 'Structura populaţiei speciei'
        self.tblA_cons.rows[16].cells[0].text = 'A.16'
        self.tblA_cons.rows[16].cells[1].text = 'Starea de conservare din punct de vedere al populaţiei speciei'
        self.tblA_cons.rows[17].cells[0].text = 'A.17'
        self.tblA_cons.rows[17].cells[1].text = 'Tendinţa stării de conservare din punct de vedere al populaţiei speciei'
        self.tblA_cons.rows[18].cells[0].text = 'A.18'
        self.tblA_cons.rows[18].cells[1].text = 'Starea de conservare necunoscută din punct de vedere al populaţiei'

    #Filler of Tabel A for species conservation
        if chk:
            self.tblA_cons.rows[1].cells[2].paragraphs[0].add_run(s.a1_lat, 'poim_lat')
            self.tblA_cons.rows[1].cells[2].paragraphs[0].add_run(s.a1)
            self.tblA_cons.rows[2].cells[2].text = s.a2
            self.tblA_cons.rows[3].cells[2].text = s.a3
            self.tblA_cons.rows[4].cells[2].text = s.a4
            self.tblA_cons.rows[5].cells[2].text = s.a5
            self.tblA_cons.rows[6].cells[2].text = s.a6
            self.tblA_cons.rows[7].cells[2].text = s.a7
            self.tblA_cons.rows[8].cells[2].text = str(s.a8)
            self.tblA_cons.rows[9].cells[2].text = s.a9
            self.tblA_cons.rows[10].cells[2].text = s.a10
            self.tblA_cons.rows[11].cells[2].text = s.a11
            self.tblA_cons.rows[12].cells[2].text = s.a12
            self.tblA_cons.rows[13].cells[2].text = s.a13
            self.tblA_cons.rows[14].cells[2].text = s.a14
            self.tblA_cons.rows[15].cells[2].text = s.a15
            self.tblA_cons.rows[16].cells[2].text = s.a16
            self.tblA_cons.rows[17].cells[2].text = s.a17
            self.tblA_cons.rows[18].cells[2].text = s.a18

    #Tabel B for species conservation
    def ftblB_cons(self, chk, s):
        ftblB_title = self.d.add_paragraph('Tabelul B) Parametri pentru evaluarea stării de conservare a speciei din punct de vedere al habitatului speciei')
        ftblB_title.style = 'poim_h2' 
        self.tblB_cons = self.d.add_table(rows=18, cols=3)
            #Set table style and column dimensions
        self.tblB_cons.style = 'poim_tg'
        for cell in self.tblB_cons.columns[0].cells:
            cell.width = Cm(2)
        for cell in self.tblB_cons.columns[1].cells:
            cell.width = Cm(5)
        for cell in self.tblB_cons.columns[2].cells:
            cell.width = Cm(9)
        self.tblB_cons.autofit = False
            #Set table fixed content
        self.tblB_cons.rows[0].cells[0].text = 'Nr.'
        self.tblB_cons.rows[0].cells[1].text = 'Parametri'
        self.tblB_cons.rows[0].cells[2].text = 'Descriere'
        self.tblB_cons.rows[1].cells[0].text = 'A.1'
        self.tblB_cons.rows[1].cells[1].text = 'Specia'
        self.tblB_cons.rows[2].cells[0].text = 'A.2'
        self.tblB_cons.rows[2].cells[1].text = 'Tipul populaţiei speciei în aria naturală protejată'
        self.tblB_cons.rows[3].cells[0].text = 'B.3'
        self.tblB_cons.rows[3].cells[1].text = 'Suprafaţa habitatului speciei în aria naturală protejată'
        self.tblB_cons.rows[4].cells[0].text = 'B.4'
        self.tblB_cons.rows[4].cells[1].text = 'Calitatea datelor pentru suprafaţa habitatului speciei'
        self.tblB_cons.rows[5].cells[0].text = 'B.5'
        self.tblB_cons.rows[5].cells[1].text = 'Suprafaţa reevaluată a habitatului speciei din planul de management anterior'
        self.tblB_cons.rows[6].cells[0].text = 'B.6'
        self.tblB_cons.rows[6].cells[1].text = 'Suprafaţa  adecvată a habitatului speciei în aria naturală protejată'
        self.tblB_cons.rows[7].cells[0].text = 'B.7'
        self.tblB_cons.rows[7].cells[1].text = 'Metodologia de apreciere a suprafeţei  adecvate a habitatului speciei în aria naturală protejată'
        self.tblB_cons.rows[8].cells[0].text = 'B.8'
        self.tblB_cons.rows[8].cells[1].text = 'Raportul dintre suprafaţa adecvată a habitatului speciei şi suprafaţa actuală a habitatului speciei'
        self.tblB_cons.rows[9].cells[0].text = 'B.9'
        self.tblB_cons.rows[9].cells[1].text = 'Tendinţa actuală a suprafeţei habitatului speciei'
        self.tblB_cons.rows[10].cells[0].text = 'B.10'
        self.tblB_cons.rows[10].cells[1].text = 'Calitatea datelor privind tendinţa actuală a suprafeţei habitatului speciei'
        self.tblB_cons.rows[11].cells[0].text = 'B.11'
        self.tblB_cons.rows[11].cells[1].text = 'Calitatea habitatului speciei în aria naturală protejată'
        self.tblB_cons.rows[12].cells[0].text = 'B.12'
        self.tblB_cons.rows[12].cells[1].text = 'Tendinţa actuală a calităţii habitatului speciei'
        self.tblB_cons.rows[13].cells[0].text = 'B.13'
        self.tblB_cons.rows[13].cells[1].text = 'Calitatea datelor privind tendinţa actuală a calităţii habitatului speciei'
        self.tblB_cons.rows[14].cells[0].text = 'B.14'
        self.tblB_cons.rows[14].cells[1].text = 'Tendinţa actuală globală a habitatului speciei funcţie de tendinţa suprafeţei şi de tendinţa calităţii habitatului speciei'
        self.tblB_cons.rows[15].cells[0].text = 'B.15'
        self.tblB_cons.rows[15].cells[1].text = 'Starea de conservare din punct de vedere al habitatului speciei'
        self.tblB_cons.rows[16].cells[0].text = 'B.16'
        self.tblB_cons.rows[16].cells[1].text = 'Tendinţa stării de conservare din punct de vedere al habitatului speciei'
        self.tblB_cons.rows[17].cells[0].text = 'B.17'
        self.tblB_cons.rows[17].cells[1].text = 'Starea de conservare necunoscută din punct de vedere al habitatului speciei'

        #Filler of Tabel B for species conservation
        if chk:
            self.tblB_cons.rows[1].cells[2].paragraphs[0].add_run(s.a1_lat, 'poim_lat')
            self.tblB_cons.rows[1].cells[2].paragraphs[0].add_run(s.a1)
            self.tblB_cons.rows[2].cells[2].text = s.a2
            self.tblB_cons.rows[3].cells[2].text = s.b3
            self.tblB_cons.rows[4].cells[2].text = s.b4
            self.tblB_cons.rows[5].cells[2].text = s.b5
            self.tblB_cons.rows[6].cells[2].text = str(s.b6) + ' ha'
            self.tblB_cons.rows[7].cells[2].text = s.b7
            self.tblB_cons.rows[8].cells[2].text = str(s.b8)
            self.tblB_cons.rows[9].cells[2].text = s.b9
            self.tblB_cons.rows[10].cells[2].text = s.b10
            self.tblB_cons.rows[11].cells[2].text = s.b11
            self.tblB_cons.rows[12].cells[2].text = s.b12
            self.tblB_cons.rows[13].cells[2].text = s.b13
            self.tblB_cons.rows[14].cells[2].text = s.b14
            self.tblB_cons.rows[15].cells[2].text = s.b15
            self.tblB_cons.rows[16].cells[2].text = s.b16
            self.tblB_cons.rows[17].cells[2].text = s.b17

    #Tabel C for species conservation
    def ftblC_cons(self, chk, s):
        ftblC_title = self.d.add_paragraph('Tabelul C) Parametri pentru evaluarea stării de conservare a speciei din punct de vedere al perspectivelor speciei în viitor')
        ftblC_title.style = 'poim_h2' 
        self.tblC_cons = self.d.add_table(rows=17, cols=3)
            #Set table style and column dimensions
        self.tblC_cons.style = 'poim_tg'
        for cell in self.tblC_cons.columns[0].cells:
            cell.width = Cm(2)
        for cell in self.tblC_cons.columns[1].cells:
            cell.width = Cm(5)
        for cell in self.tblC_cons.columns[2].cells:
            cell.width = Cm(9)
        self.tblC_cons.autofit = False
            #Set table fixed content
        self.tblC_cons.rows[0].cells[0].text = 'Nr.'
        self.tblC_cons.rows[0].cells[1].text = 'Parametru'
        self.tblC_cons.rows[0].cells[2].text = 'Descriere'
        self.tblC_cons.rows[1].cells[0].text = 'A.1'
        self.tblC_cons.rows[1].cells[1].text = 'Specia'
        self.tblC_cons.rows[2].cells[0].text = 'A.2'
        self.tblC_cons.rows[2].cells[1].text = 'Tipul populaţiei speciei în aria naturală protejată'
        self.tblC_cons.rows[3].cells[0].text = 'C.3'
        self.tblC_cons.rows[3].cells[1].text = 'Tendinţa viitoare a mărimii populaţiei'
        self.tblC_cons.rows[4].cells[0].text = 'C.4'
        self.tblC_cons.rows[4].cells[1].text = 'Raportul dintre mărimea populaţiei de referinţă pentru starea favorabilă şi mărimea populaţiei viitoare a speciei'
        self.tblC_cons.rows[5].cells[0].text = 'C.5'
        self.tblC_cons.rows[5].cells[1].text = 'Perspectivele speciei din punct de vedere al populaţiei'
        self.tblC_cons.rows[6].cells[0].text = 'C.6'
        self.tblC_cons.rows[6].cells[1].text = 'Tendinţa viitoare a suprafeţei habitatului speciei'
        self.tblC_cons.rows[7].cells[0].text = 'C.7'
        self.tblC_cons.rows[7].cells[1].text = 'Raportul dintre suprafaţa adecvată a habitatului speciei şi suprafaţa habitatului speciei în viitor'
        self.tblC_cons.rows[8].cells[0].text = 'C.8'
        self.tblC_cons.rows[8].cells[1].text = 'Perspectivele speciei din punct de vedere al habitatului speciei'
        self.tblC_cons.rows[9].cells[0].text = 'C.9'
        self.tblC_cons.rows[9].cells[1].text = 'Perspectivele speciei în viitor'
        self.tblC_cons.rows[10].cells[0].text = 'C.10'
        self.tblC_cons.rows[10].cells[1].text = 'Efectul cumulat al impacturilor asupra speciei în viitor'
        self.tblC_cons.rows[11].cells[0].text = 'C.11'
        self.tblC_cons.rows[11].cells[1].text = 'Intensitatea presiunilor actuale asupra speciei'
        self.tblC_cons.rows[12].cells[0].text = 'C.12'
        self.tblC_cons.rows[12].cells[1].text = 'Intensitatea ameninţărilor viitoare asupra speciei'
        self.tblC_cons.rows[13].cells[0].text = 'C.13'
        self.tblC_cons.rows[13].cells[1].text = 'Viabilitatea pe termen lung a speciei'
        self.tblC_cons.rows[14].cells[0].text = 'C.14'
        self.tblC_cons.rows[14].cells[1].text = 'Starea de conservare din punct de vedere al perspectivelor speciei în viitor'
        self.tblC_cons.rows[15].cells[0].text = 'C.15'
        self.tblC_cons.rows[15].cells[1].text = 'Tendinţa stării de conservare din punct de vedere al perspectivelor speciei în viitor'
        self.tblC_cons.rows[16].cells[0].text = 'C.16'
        self.tblC_cons.rows[16].cells[1].text = 'Starea de conservare necunoscută din punct de vedere al perspectivelor speciei în viitor'


        #Filler of Tabel C for species conservation
        if chk:
            self.tblC_cons.rows[1].cells[2].paragraphs[0].add_run(s.a1_lat, 'poim_lat')
            self.tblC_cons.rows[1].cells[2].paragraphs[0].add_run(s.a1)
            self.tblC_cons.rows[2].cells[2].text = s.a2
            self.tblC_cons.rows[3].cells[2].text = s.c3
            self.tblC_cons.rows[4].cells[2].text = s.c4
            self.tblC_cons.rows[5].cells[2].text = s.c5
            self.tblC_cons.rows[6].cells[2].text = s.c6
            self.tblC_cons.rows[7].cells[2].text = s.c7
            self.tblC_cons.rows[8].cells[2].text = s.c8
            self.tblC_cons.rows[9].cells[2].text = s.c9
            self.tblC_cons.rows[10].cells[2].text = s.c10
            self.tblC_cons.rows[11].cells[2].add_paragraph('Ridicat: ', 'poim_th')
            if len(s.c11r)>0:
                for x in range(len(s.c11r)):
                    self.tblC_cons.rows[11].cells[2].add_paragraph(s.c11r[x])
            else:
                self.tblC_cons.rows[11].cells[2].add_paragraph('Pentru această specie nu se cunosc presiuni cu intensitate ridicată')
            self.tblC_cons.rows[11].cells[2].add_paragraph('Mediu: ', 'poim_th')
            if len(s.c11m)>0:
                for x in range(len(s.c11m)):
                    self.tblC_cons.rows[11].cells[2].add_paragraph(s.c11m[x])
            else:
                self.tblC_cons.rows[11].cells[2].add_paragraph('Pentru această specie nu se cunosc presiuni cu intensitate medie')
            self.tblC_cons.rows[11].cells[2].add_paragraph('Scăzut: ', 'poim_th')
            if len(s.c11s)>0:
                for x in range(len(s.c11s)):
                    self.tblC_cons.rows[11].cells[2].add_paragraph(s.c11s[x])
            else:
                self.tblC_cons.rows[11].cells[2].add_paragraph('Pentru această specie nu se cunosc presiuni cu intensitate scăzută')
            
            
            self.tblC_cons.rows[12].cells[2].add_paragraph('Ridicat: ', 'poim_th')

            if len(s.c12r)>0:
                for x in range(len(s.c12r)):
                    self.tblC_cons.rows[12].cells[2].add_paragraph(s.c12r[x])
            else:
                self.tblC_cons.rows[12].cells[2].add_paragraph('Pentru această specie nu se cunosc amenințări cu intensitate ridicată')
            self.tblC_cons.rows[12].cells[2].add_paragraph('Mediu: ', 'poim_th')
            if len(s.c12m)>0:
                for x in range(len(s.c12m)):
                    self.tblC_cons.rows[12].cells[2].add_paragraph(s.c12m[x])
            else:
                self.tblC_cons.rows[12].cells[2].add_paragraph('Pentru această specie nu se cunosc amenințări cu intensitate medie')
            self.tblC_cons.rows[12].cells[2].add_paragraph('Scăzut: ', 'poim_th')
            if len(s.c12s)>0:
                for x in range(len(s.c12s)):
                    self.tblC_cons.rows[12].cells[2].add_paragraph(s.c12s[x])
            else:
                self.tblC_cons.rows[12].cells[2].add_paragraph('Pentru această specie nu se cunosc amenințări cu intensitate scăzută')

            self.tblC_cons.rows[13].cells[2].text = s.c13
            self.tblC_cons.rows[14].cells[2].text = s.c14
            self.tblC_cons.rows[15].cells[2].text = s.c15
            self.tblC_cons.rows[16].cells[2].text = s.c16

    #Tabel D for species conservation
    def ftblD_cons(self, chk, s):
        ftblD_title = self.d.add_paragraph('Tabelul D) Parametri pentru evaluarea stării globale de conservare a speciei în cadrul ariei naturale protejate')
        ftblD_title.style = 'poim_h2' 
        self.tblD_cons = self.d.add_table(rows=7, cols=3)
        #Set table style and column dimensions
        self.tblD_cons.style = 'poim_tg'
        for cell in self.tblD_cons.columns[0].cells:
            cell.width = Cm(2)
        for cell in self.tblD_cons.columns[1].cells:
            cell.width = Cm(5)
        for cell in self.tblD_cons.columns[2].cells:
            cell.width = Cm(9)
        self.tblD_cons.autofit = False
        #Set table fixed content
        self.tblD_cons.rows[0].cells[0].text = 'Nr.'
        self.tblD_cons.rows[0].cells[1].text = 'Parametru'
        self.tblD_cons.rows[0].cells[2].text = 'Descriere'
        self.tblD_cons.rows[1].cells[0].text = 'A.1'
        self.tblD_cons.rows[1].cells[1].text = 'Specia'
        self.tblD_cons.rows[2].cells[0].text = 'A.2'
        self.tblD_cons.rows[2].cells[1].text = 'Tipul populaţiei speciei în aria naturală protejată'
        self.tblD_cons.rows[3].cells[0].text = 'D.3'
        self.tblD_cons.rows[3].cells[1].text = 'Starea globală de conservare a speciei'
        self.tblD_cons.rows[4].cells[0].text = 'D.4'
        self.tblD_cons.rows[4].cells[1].text = 'Tendinţa stării globale de conservare a speciei'
        self.tblD_cons.rows[5].cells[0].text = 'D.5'
        self.tblD_cons.rows[5].cells[1].text = 'Starea globală de conservare necunoscută'
        self.tblD_cons.rows[6].cells[0].text = 'D.6'
        self.tblD_cons.rows[6].cells[1].text = 'Informaţii suplimentare'    

        #Filler of Tabel D for species conservation
        if chk:
            self.tblD_cons.rows[1].cells[2].paragraphs[0].add_run(s.a1_lat, 'poim_lat')
            self.tblD_cons.rows[1].cells[2].paragraphs[0].add_run(s.a1)
            self.tblD_cons.rows[2].cells[2].text = s.a2
            self.tblD_cons.rows[3].cells[2].text = s.d3
            self.tblD_cons.rows[4].cells[2].text = s.d4
            self.tblD_cons.rows[5].cells[2].text = s.d5
            self.tblD_cons.rows[6].cells[2].text = s.d6

    #Tabel for matrix 1
    def ftblM_1(self, chk, s):
        ftblM_1_title = self.d.add_paragraph('Matricea 1) Matricea de evaluare a stării de conservare a speciei din punct de vedere al populației speciei')
        ftblM_1_title.style = 'poim_h2' 
        self.tblM_1 = self.d.add_table(rows=2, cols=4)
        self.tblM_1.style = 'poim_tg'
        for cell in self.tblM_1.columns[0].cells:
            cell.width = Cm(4)
        for cell in self.tblM_1.columns[1].cells:
            cell.width = Cm(4)
        for cell in self.tblM_1.columns[2].cells:
            cell.width = Cm(4)
        for cell in self.tblM_1.columns[3].cells:
            cell.width = Cm(4)
        self.tblM_1.autofit = False
        #Set table fixed content
        self.tblM_1.rows[0].cells[0].text = 'Favorabilă'
        self.tblM_1.rows[0].cells[1].text = 'Nefavorabilă -Inadecvată'
        self.tblM_1.rows[0].cells[2].text = 'Nefavorabilă - Rea'
        self.tblM_1.rows[0].cells[3].text = 'Necunoscută'

        #Filler of Tabel for matrix 1
        if chk:
            self.tblM_1.rows[1].cells[s.col_cons_pop].text = str(s.txt_cons_pop)

    #Tabel for matrix 2
    def ftblM_2(self, chk, s):
        ftblM_2_title = self.d.add_paragraph('Matricea 2) Matricea pentru evaluarea tendinței globale a habitatului speciei')
        ftblM_2_title.style = 'poim_h2' 
        self.tblM_2 = self.d.add_table(rows=2, cols=2)
        self.tblM_2.style = 'poim_tg'
        for cell in self.tblM_2.columns[0].cells:
            cell.width = Cm(8)
        for cell in self.tblM_2.columns[1].cells:
            cell.width = Cm(8)
        self.tblM_2.autofit = False
            #Set table fixed content
        self.tblM_2.rows[0].cells[0].text = 'Tendinţa'
        self.tblM_2.rows[0].cells[1].text = 'Combinaţia dintre Tendinţa actuală a suprafeţei habitatului speciei [B.9.] şi Tendinţa actuală a calităţii habitatului speciei [B.12.]'

        #Filler of Tabel for matrix 2
        if chk:
            self.tblM_2.rows[1].cells[0].text = s.b14
            self.tblM_2.rows[1].cells[1].text = 'B9: '+ s.b9 + ', B12: '+ s.b12

    #Tabel for matrix 3
    def ftblM_3(self, chk, s):
        ftblM_3_title = self.d.add_paragraph('Matricea 3) Matricea de evaluare a stării de conservare a speciei din punct de vedere al habitatului speciei')
        ftblM_3_title.style = 'poim_h2' 
        self.tblM_3 = self.d.add_table(rows=2, cols=4)
        self.tblM_3.style = 'poim_tg'
        for cell in self.tblM_3.columns[0].cells:
            cell.width = Cm(4)
        for cell in self.tblM_3.columns[1].cells:
            cell.width = Cm(4)
        for cell in self.tblM_3.columns[2].cells:
            cell.width = Cm(4)
        for cell in self.tblM_3.columns[3].cells:
            cell.width = Cm(4)
        self.tblM_3.autofit = False
            #Set table fixed content
        self.tblM_3.rows[0].cells[0].text = 'Favorabilă'
        self.tblM_3.rows[0].cells[1].text = 'Nefavorabilă -Inadecvată'
        self.tblM_3.rows[0].cells[2].text = 'Nefavorabilă - Rea'
        self.tblM_3.rows[0].cells[3].text = 'Necunoscută'

        #Filler of Tabel for matrix 3
        if chk:
            self.tblM_3.rows[1].cells[s.col_cons_hab].text = 'Raportul dintre suprafaţa adecvată a habitatului speciei şi suprafaţa actuală a habitatului speciei: ' + s.b88 + '. Tendinţa actuală a suprafeţei habitatului speciei: ' + s.b9 + '. Calitatea habitatului speciei în aria naturală protejată: ' + s.b11

    #Tabel for matrix 4
    def ftblM_4(self, chk, s):
        ftblM_4_title = self.d.add_paragraph('Matricea 4) Matricea pentru evaluarea perspectivelor speciei din punct de vedere al populației speciei')
        ftblM_4_title.style = 'poim_h2' 
        self.tblM_4 = self.d.add_table(rows=2, cols=5)
        self.tblM_4.style = 'poim_tg'
        for cell in self.tblM_4.columns[0].cells:
            cell.width = Cm(3)
        for cell in self.tblM_4.columns[1].cells:
            cell.width = Cm(3)
        for cell in self.tblM_4.columns[2].cells:
            cell.width = Cm(3)
        for cell in self.tblM_4.columns[3].cells:
            cell.width = Cm(3)
        for cell in self.tblM_4.columns[4].cells:
            cell.width = Cm(4)   
        self.tblM_4.autofit = False
            #Set table fixed content
        self.tblM_4.rows[0].cells[0].text = 'Valoarea actuală a parametrului'
        self.tblM_4.rows[0].cells[1].text = 'Tendinţa viitoare a parametrului'
        self.tblM_4.rows[0].cells[2].text = 'Raportul dintre valoarea VRSF şi valoarea viitoare a parametrului'
        self.tblM_4.rows[0].cells[3].text = 'Perspective'
        self.tblM_4.rows[0].cells[4].text = 'Figura'

        #Filler of Tabel for matrix 4
        if chk:
            self.tblM_4.rows[1].cells[0].text = s.a10
            self.tblM_4.rows[1].cells[1].text = s.c3
            self.tblM_4.rows[1].cells[2].text = s.c4
            self.tblM_4.rows[1].cells[3].text = s.c5
            run_for_image = self.tblM_4.rows[1].cells[4].paragraphs[0].add_run()
            if s.mat4_14 == 'na':
                self.tblM_4.rows[1].cells[4].text = 'Nu exista imagine corespunzatoare acestei situatii'
            else:
                run_for_image.add_picture(s.mat4_14, width=Cm(3.5))
            
    
    #Tabel for matrix 5
    def ftblM_5(self, chk, s):
        ftblM_5_title = self.d.add_paragraph('Matricea 5) Perspectivele speciei în viitor, după implementarea planului de management actual')
        ftblM_5_title.style = 'poim_h2' 
        self.tblM_5 = self.d.add_table(rows=2, cols=4)
        self.tblM_5.style = 'poim_tg'
        for cell in self.tblM_5.columns[0].cells:
            cell.width = Cm(4)
        for cell in self.tblM_5.columns[1].cells:
            cell.width = Cm(4)
        for cell in self.tblM_5.columns[2].cells:
            cell.width = Cm(4)
        for cell in self.tblM_5.columns[3].cells:
            cell.width = Cm(4)
        self.tblM_5.autofit = False
            #Set table fixed content
        self.tblM_5.rows[0].cells[0].text = 'Favorabilă'
        self.tblM_5.rows[0].cells[1].text = 'Nefavorabilă -Inadecvată'
        self.tblM_5.rows[0].cells[2].text = 'Nefavorabilă - Rea'
        self.tblM_5.rows[0].cells[3].text = 'Necunoscută'

        #Filler of Tabel for matrix 5
        if chk:
            self.tblM_5.rows[1].cells[s.col_cons_pre].text = 'Perspectivele speciei din punct de vedere al populaţiei: ' + s.c5 + '. Perspectivele speciei din punct de vedere al habitatului speciei: ' + s.c8

    #Tabel for matrix 6
    def ftblM_6(self, chk, s):
        ftblM_6_title = self.d.add_paragraph('Matricea 6) Matricea evaluării stării de conservare a speciei din punct de vedere al perspectivelor speciei în viitor, după implementarea planului de management actual')
        ftblM_6_title.style = 'poim_h2' 
        self.tblM_6 = self.d.add_table(rows=2, cols=4)
        self.tblM_6.style = 'poim_tg'
        for cell in self.tblM_6.columns[0].cells:
            cell.width = Cm(4)
        for cell in self.tblM_6.columns[1].cells:
            cell.width = Cm(4)
        for cell in self.tblM_6.columns[2].cells:
            cell.width = Cm(4)
        for cell in self.tblM_6.columns[3].cells:
            cell.width = Cm(4)
        self.tblM_6.autofit = False
            #Set table fixed content
        self.tblM_6.rows[0].cells[0].text = 'Favorabilă'
        self.tblM_6.rows[0].cells[1].text = 'Nefavorabilă -Inadecvată'
        self.tblM_6.rows[0].cells[2].text = 'Nefavorabilă - Rea'
        self.tblM_6.rows[0].cells[3].text = 'Necunoscută'

        #Filler of Tabel for matrix 6
        if chk:
            self.tblM_6.rows[1].cells[s.mat6_col].text = s.txt_mat6_col     

    #Tabel for matrix 7
    def ftblM_7(self, chk, s):
        ftblM_7_title = self.d.add_paragraph('Matricea 7) Evaluarea stării globale de conservare a speciei')
        ftblM_7_title.style = 'poim_h2' 
        self.tblM_7 = self.d.add_table(rows=2, cols=4)
        self.tblM_7.style = 'poim_tg'
        for cell in self.tblM_7.columns[0].cells:
            cell.width = Cm(4)
        for cell in self.tblM_7.columns[1].cells:
            cell.width = Cm(4)
        for cell in self.tblM_7.columns[2].cells:
            cell.width = Cm(4)
        for cell in self.tblM_7.columns[3].cells:
            cell.width = Cm(4)
        self.tblM_7.autofit = False
            #Set table fixed content
        self.tblM_7.rows[0].cells[0].text = 'Favorabilă'
        self.tblM_7.rows[0].cells[1].text = 'Nefavorabilă -Inadecvată'
        self.tblM_7.rows[0].cells[2].text = 'Nefavorabilă - Rea'
        self.tblM_7.rows[0].cells[3].text = 'Necunoscută'

        #Filler of Tabel for matrix 7
        if chk:
            self.tblM_7.rows[1].cells[s.col_cons_glob].text = s.txt_cons_glob



    #SINTETIC TABLE 
    def sintetic_table_head(self):
        self.tb_sin = self.d.add_table(rows=1, cols=4)
        self.tb_sin.style = 'poim_tg'
        for cell in self.tb_sin.columns[0].cells:
            cell.width = Cm(3)
        for cell in self.tb_sin.columns[1].cells:
            cell.width = Cm(2)
        for cell in self.tb_sin.columns[2].cells:
            cell.width = Cm(6.5)
        for cell in self.tb_sin.columns[3].cells:
            cell.width = Cm(5)
        self.tb_sin.autofit = False    
        self.tb_sin.rows[0].cells[0].text = 'Aria naturală protejată / Elementele de interes conservativ'
        self.tb_sin.rows[0].cells[1].text = 'Starea de conservare (F/NI/NR)'
        self.tb_sin.rows[0].cells[2].text = 'Presiune (P)  / Amenințare (A) (cod)'
        self.tb_sin.rows[0].cells[3].text =  'Măsuri de conservare propuse'

    def sintetic_table_row(self, s):
        
        self.tb_sin = self.d.add_table(rows=1, cols=4)
        self.tb_sin.style = 'poim_tg'
        for cell in self.tb_sin.columns[0].cells:
            cell.width = Cm(3)
        for cell in self.tb_sin.columns[1].cells:
            cell.width = Cm(2)
        for cell in self.tb_sin.columns[2].cells:
            cell.width = Cm(6.5)
        for cell in self.tb_sin.columns[3].cells:
            cell.width = Cm(5)
        self.tb_sin.autofit = False
            #Set table content
        self.tb_sin.rows[0].cells[0].add_paragraph(s.lat_sp + ' - '+s.feno)
        self.tb_sin.rows[0].cells[1].add_paragraph(s.d3_sin)
        if len(s.c11_sin)>0:
            for x in s.c11_sin:        
                self.tb_sin.rows[0].cells[2].add_paragraph('P - '+ x)
        if len(s.c11_sin)>0:
            for x in s.c11_sin:        
                self.tb_sin.rows[0].cells[2].add_paragraph('A - '+ x)
        if len(s.ms_list)>0:
            for x in s.ms_list:        
                self.tb_sin.rows[0].cells[3].add_paragraph(x)

            

    #Redefine save method
    def save(self, fn):
        self.d.save(fn)





class specie:
    '''
    Species object designed to work with pm_doc object.
    Takes attributes from 5 dataframes: 
    df = google sheet with all tabel cell filled by experts , 
    bf = tabel from theone database with different info
    pf = google sheet with species descriptions 
    impf = google sheet cu impacturile asupra speciilor
    msf = google sheet cu masurile propuse
    '''



    def __init__(self, row, df, bf, pf, impf, msf):
        self.row = row


                                                    #Seteaza dictionarele necesare pentru calcularea atributelor
            # Variabilele de tip text folosite pt fenologie
        feno_dict={
            'C':'Populaţie aflată în pasaj care utilizează aria naturală protejată pentru odihnă şi/sau hrănire.', 
            'P':'Populaţie permanentă (sedentară/rezidentă)', 
            'R':'Populaţie nerezidentă cuibăritoare (care utilizează aria naturală protejată pentru reproducere)',
            'W':'Populaţie care doar iernează în aria naturală protejată',
            'na': 'na'
        }

            # Variabilele de tip text folosite pt starea de conservare
        con={
            'FV':'”FV” – favorabilă',
            'U1':'”U1” – nefavorabilă - inadecvată',
            'U2':'”U2” – nefavorabilă - rea',
            'X':'”X” – necunoscută',
            'na': 'na'
        }
            # Variabilele de tip text folosite pentru calitatea datelor
        cal={
            3:'bună - estimări statistice robuste sau inventarieri complete',
            2:'medie - date estimate pe baza extrapolării şi/sau modelării datelor obţinute prin măsurători parţiale',
            1:'slabă - date estimate pe baza opiniei experţilor cu sau fără măsurători prin eşantionare',
            0:'insuficientă – date insuficiente sau nesigure',
            'na': 'na'
        }
            # Variabilele text folosite la raporturile dintre populatii
        rap={
            1:'”<” – mai mic (în condiţii excepţionale)',
            2:'”≈” – aproximativ egal',
            3:'”>” – mai mare',
            4:'”>>” – mult mai mare',
            0:'”x” – necunoscut',
            'na': 'na'
        }
            # Variabilele text folosite pentru a desemna tendintele unor populatii, a suprafetelor, a calitatii
        ten={
            3:'”+” –crescătoare',
            1:'”-” – descrescătoare',
            2:'”0” – stabilă',
            0:'”x” – necunoscută',
            'na': 'na'
        }
            # Variabilele text folosite pentru a desemna tendintele unor stari
        ten_con={
            'c':'”+” – se îmbunătăţeşte',
            'd':'”-” – se înrăutăţeşte',
            's':'”0” – este stabilă',
            'n':'”x” – necunoscută',
            'na': 'na'
        }
            #Variabilele privind perspectivele (C5, C8)
        persp={
            'FV':'”FV” –  favorabile',
            'U1':'”U1” –  nefavorabile - inadecvate',
            'U2':'”U2” –  nefavorabile - rele',
            'X':'”X” –  necunoscute',
            'na': 'na'
        }
        imp={
            'r':'Ridicat - impacturile‚ respectiv presiunile actuale și/sau amenințările viitoare vor avea în viitor un efect cumulat ridicat asupra speciei‚ afectând major viabilitatea pe termen lung a speciei',
            'm':'Mediu - impacturile‚ respectiv presiunile actuale și/sau amenințările viitoare vor avea în viitor un efect cumulat mediu‚ semnificativ asupra speciei‚ afectând semnificativ viabilitatea pe termen lung a speciei',
            's':'Scăzut - impacturile‚ respectiv presiunile actuale și amenințările viitoare‚ vor avea un efect cumulat scăzut sau nesemnificativ asupra speciei‚ neafectând semnificativ viabilitatea pe termen lung a speciei',
            'n':'Nu există suficiente informații în ceea ce privește efectul impacturilor asupra speciei în viitor',
            'na': 'na'
        }


    #Codul bl9_id din baza de date (df)
        self.cod_sp = df.loc[row]['id']
        logging.debug('A setat cod_sp'+' = ' +str(self.cod_sp))
    #Fenologia speciei (df)
        self.feno = df.loc[row]['feno']
        logging.debug('A setat feno'+' = ' +str(self.feno))
    #Slicing the dataframe to get onli what we want
        cd1 = bf['bl9_id'] == int(float(self.cod_sp))
        cd2 = bf['feno'] == self.feno
        sp_bf = bf[cd1&cd2]
        logging.debug('A setat sp_bf like this')
        logging.debug(sp_bf)
    #Codul Natura_2000 din baza de date (df)
        #self.cod_n = bf[bf['bl9_id'] == int(float(self.cod_sp)][bf['feno'] == self.feno]['codn'].item()
        self.cod_n = next(iter(sp_bf['codn']), 'na')
        #self.cod_n = next(iter(bf[bf['bl9_id'] == int(float(self.cod_sp))][bf['feno'] == self.feno]['codn']), 'na')
        logging.debug('A setat cod_n'+' = ' +str(self.cod_n))
    #Denumirea in latina a speciei (df)
        self.lat_sp = df.loc[row]['den_latin']
        logging.debug('A setat lat_sp'+' = ' +str(self.lat_sp))
    #Denumirea in romana a speciei (bf)
        #self.rom_sp = next(iter(bf[bf['bl9_id'] == int(float(self.cod_sp))][bf['feno'] == self.feno]['denro']), 'na')
        self.rom_sp = next(iter(sp_bf['denro']), 'na')
        logging.debug('A setat rom_sp'+' = ' +str(self.rom_sp))
    #Codul eunis al speciei  (bf)
        #self.eunis = next(iter(bf[bf['bl9_id'] == int(float(self.cod_sp))][bf['feno'] == self.feno]['idspecie']), 'na')
        self.eunis = next(iter(sp_bf['idspecie']), 'na')
        logging.debug('A setat eunis'+' = ' +str(self.eunis))       
    #Anexa DP pe care este specia (db)
        #self.anexa = next(iter(bf[bf['bl9_id'] == int(float(self.cod_sp))][bf['feno'] == self.feno]['nbanexa']), 'na')
        self.anexa = next(iter(sp_bf['nbanexa']), 'na')
        logging.debug('A setat anexa'+' = ' +str(self.anexa))
    #Populatia minima in sit la acest pm(df)
        self.pop_sit_min = int(float(df.loc[row]['a3_min'])) if df.loc[row]['a3_min'] != 'na' else 'na'
        logging.debug('A setat pop_min_sit'+' = ' +str(self.pop_sit_min))
    #Populatia maxima in sit la acest pm (df)
        self.pop_sit_max = int(float(df.loc[row]['a3_max'])) if df.loc[row]['a3_max'] != 'na' else  'na'
        logging.debug('A setat pop_max_sit'+' = ' +str(self.pop_sit_max))
    #Populatia minima in sit pe formularul standard(df)
        #if next(iter(bf[(bf['bl9_id'] == int(float(self.cod_sp))) & (bf['feno'] == self.feno)]['popmin']),'na') == 'na':
        if next(iter(sp_bf['popmin']), 'na') == 'na':
            self.pop_for_min = 'na'
        else:   
            self.pop_for_min = int(float(next(iter(sp_bf['popmin']), 'na')))
        logging.debug('A setat pop_for_min'+' = ' +str(self.pop_for_min))
    #Populatia maxima in sit pe formularul standard (df)
        if next(iter(sp_bf['popmax']),'na') == 'na':
            self.pop_for_max = 'na'
        else:   
            self.pop_for_max = int(float(next(iter(sp_bf['popmax']),'na')))
        logging.debug('A setat pop_for_max'+' = ' +str(self.pop_for_max))
    #Populatia nationala minima (db)
        if next(iter(sp_bf['popnmin']),'na') == 'na':
            self.pop_nat_min = 'na'
        else:
            self.pop_nat_min = int(float(next(iter(sp_bf['popnmin']),'na')))
        logging.debug('A setat pop_nat_min'+' = ' +str(self.pop_nat_min))
    #Populatia nationala maxima (db)
        if next(iter(sp_bf['popnmax']),'na') == 'na':
            self.pop_nat_max = 'na'
        else:
            self.pop_nat_max = int(float(next(iter(sp_bf['popnmax']),'na')))
        logging.debug('A setat pop_nat_max'+' = ' +str(self.pop_nat_max))
    #Unitatea folosita pentru specie (alg)
        self.unitate = 'perechi' if self.feno in ['R', 'P'] else 'indivizi'
        logging.debug('A setat unitatea'+' = ' +str(self.unitate))
    #Populatia nationala medie (alg)
        self.pop_nat = round((self.pop_nat_min + self.pop_nat_max)/2) if self.pop_nat_min != 'na' and self.pop_nat_max != 'na'  else 'na'
        logging.debug('A setat pop_nat'+' = ' +str(self.pop_nat))
    #Populatia medie in sit (alg)
        self.pop_sit = math.ceil((self.pop_sit_min + self.pop_sit_max)/2) if self.pop_sit_min != 'na' and self.pop_sit_max != 'na' else 'na'
        logging.debug('A setat pop_sit'+' = ' +str(self.pop_sit))
    #Populatia medie de pe formularul standard (alg + db)
        self.pop_spa = round((self.pop_for_min + self.pop_for_max)/2) if self.pop_for_min != 'na' and self.pop_for_max != 'na' else 'na'
        logging.debug('A setat pop_spa'+' = ' +str(self.pop_spa))        


    #Attributes for Tabel A for species description
        self.as1 = str(self.eunis)
        logging.debug('A setat self.as1'+' = ' +str(self.as1))

        self.as2 = self.lat_sp
        logging.debug('A setat self.as2'+' = ' +str(self.as2))

        self.as3 = self.rom_sp
        logging.debug('A setat self.as3'+' = ' +str(self.as3))

        #self.as4 = 'Descrierea generala unitara'
        #logging.debug('A setat self.as4'+' = ' +str(self.as4))

        self.as4_gen = next(iter(pf[pf['bl9_id'] == int(float(self.cod_sp))]['desc_sp']),'na')
        logging.debug('A setat self.as4_gen'+' = ' +str(self.as4_gen))

        self.as4_hrana = next(iter(pf[pf['bl9_id'] == int(float(self.cod_sp))]['desc_hrana']),'na')
        logging.debug('A setat self.as4_hrana'+' = ' +str(self.as4_hrana))

        self.as4_cuib = next(iter(pf[pf['bl9_id'] == int(float(self.cod_sp))]['desc_cuibarit']),'na')
        logging.debug('A setat self.as4_cuib'+' = ' +str(self.as4_cuib))

        if self.feno == 'C':
            self.as5 = next(iter(pf[pf['bl9_id'] == int(float(self.cod_sp))]['per_mig']),'na')
        elif self.feno in ['R', 'P']:
            self.as5 = next(iter(pf[pf['bl9_id'] == int(float(self.cod_sp))]['per_cuib']),'na')
        elif self.feno == 'W':
            self.as5 = next(iter(pf[pf['bl9_id'] == int(float(self.cod_sp))]['per_win']),'na')
        else:
            self.as5 = 'Nu au fost identificate perioade critice.'
        logging.debug('A setat self.as5'+' = ' +str(self.as5))

        self.as6 = next(iter(pf[pf['bl9_id'] == int(float(self.cod_sp))]['desc_habitat']),'na')
        logging.debug('A setat self.as6'+' = ' +str(self.as6))

        if glob.glob('img/' + str(self.cod_sp) +'*.jpg'):
            self.as7 = glob.glob('img/' + str(self.cod_sp) +'*.jpg')[0]
        else:
            self.as7 = 'na'
        logging.debug('A setat self.as7'+' = ' +str(self.as7))

    #Attributes for Tabel B for species description
        self.bs1 = self.lat_sp 
        self.bs11 = ', ' + str(self.eunis) + ', ' + self.anexa
        logging.debug('A setat self.bs1'+' = ' +str(self.bs1))

        self.bs2 = df.loc[row]['info_spec']
        logging.debug('A setat self.bs2'+' = ' +str(self.bs2))

        if self.feno == 'C':
            self.bs3 = 'odihnă şi hranire / pasaj'
        elif self.feno == 'R':
            self.bs3 = 'reproducere'
        elif self.feno == 'P':
            self.bs3 = 'rezident'
        elif self.feno == 'W':
            self.bs3 = 'iernare'
        else:
            self.bs3 = 'na'        
        logging.debug('A setat self.bs3'+' = ' +str(self.bs3))

        self.bs4 = df.loc[row]['prez_spt']
        logging.debug('A setat self.bs4'+' = ' +str(self.bs4))

        self.bs5 = df.loc[row]['prez_mng']
        logging.debug('A setat self.bs5'+' = ' +str(self.bs5))

        self.bs6 = df.loc[row]['ab']
        logging.debug('A setat self.bs6'+' = ' +str(self.bs6))

        self.bs7 = df.loc[row]['colect']
        logging.debug('A setat self.bs7'+' = ' +str(self.bs7))

        self.bs8 = df.loc[row]['info_dist']
        logging.debug('A setat self.bs8'+' = ' +str(self.bs8))

        if glob.glob('map/*' + self.cod_n +'-'+ self.feno +'.jp*'):
            self.bs9 = glob.glob('map/*' + self.cod_n +'-'+ self.feno +'.jp*')[0]
        else:
            self.bs9 = 'na'
        logging.debug('A setat self.bs9'+' = ' +str(self.bs9))

        self.bs10 = df.loc[row]['info']
        logging.debug('A setat self.bs10'+' = ' +str(self.bs10))



        #Othter required attributes
        self.b4 = cal[df.loc[row]['b4']].split(' - ')[0]
        logging.debug('A setat self.b4'+' = ' +str(self.b4))




        #Attributes for Tabel A for species conservation
    #A.1.	Specia (alg)
        self.a1 =  ', ' + str(self.eunis) + ', ' + self.anexa
        self.a1_lat = self.lat_sp
        logging.debug('A setat A1'+' = ' + str(self.a1_lat) + ' ' +str(self.a1))
    #A.2	Statut de prezenţă temporală a speciilor (alg)
        self.a2 = feno_dict[self.feno]
        logging.debug('A setat A2'+' = ' +str(self.a2))
    #A.3	Mărimea populaţiei speciei în aria naturală protejată (alg)
        self.a3 = str(self.pop_sit_min) +' - '+ str(self.pop_sit_max) + ' ' + self.unitate
        logging.debug('A setat A3'+' = ' +str(self.a3))
    #A.4	Calitatea datelor referitoare la populaţia speciei din aria naturală protejată (df)
        self.a4 = cal[df.loc[row]['a4']]
        logging.debug('A setat A4'+' = ' +str(self.a4))
    #A.5	Raportul dintre mărimea populaţiei speciei în aria naturală protejată şi mărimea populaţiei naţionale (alg)
        if self.pop_nat_min == 0:
            self.a5 = str(self.pop_sit_min) + '/' + str(self.pop_nat_min)+ ' %' + ' - '+ str(int(self.pop_sit_max*100/self.pop_nat_max))+ ' %' if self.pop_nat_min != 'na' and self.pop_nat_max != 'na'  else 'Nu exista date disponibile pentru stabilirea acestui parametru'
            if self.pop_nat_max == 0:
                self.a5 = str(int(self.pop_sit_min*100/self.pop_nat_min))+ ' %' + ' - '+ str(self.pop_sit_max) + '/' + str(self.pop_nat_max)+ ' %' if self.pop_nat_min != 'na' and self.pop_nat_max != 'na'  else 'Nu exista date disponibile pentru stabilirea acestui parametru'
        elif self.pop_nat_max == 0:
            self.a5 = str(int(self.pop_sit_min*100/self.pop_nat_min))+ ' %' + ' - '+ str(self.pop_sit_max) + '/' + str(self.pop_nat_max)+ ' %' if self.pop_nat_min != 'na' and self.pop_nat_max != 'na'  else 'Nu exista date disponibile pentru stabilirea acestui parametru'       
        else:           
            self.a5 = str(round(self.pop_sit_min*100/self.pop_nat_min,2))+ ' %' + ' - '+ str(round(self.pop_sit_max*100/self.pop_nat_max,2))+ ' %' if self.pop_nat_min != 'na' and self.pop_nat_max != 'na'  else 'Nu exista date disponibile pentru stabilirea acestui parametru'
        logging.debug('A setat A5'+' = ' +str(self.a5))
    #A.6	Mărimea populaţiei speciei în aria naturală protejată comparata cu mărimea populaţiei naţionale (alg)
        self.a6 = 'Semnificativă' if self.pop_nat_min != 'na' and self.pop_nat_max != 'na' and self.pop_sit/self.pop_nat >= 0.1 else 'Nesemnificativă'
        logging.debug('A setat A6'+' = ' + self.a6)
    #A.7	Mărimea reevaluată a populaţiei estimate în planul de management anterior (cst)
        if df.loc[row]['a7_min'] != 'na':
            self.a7_min = int(float(df.loc[row]['a7_min']))
        else:
            self.a7_min = 'na'

        if df.loc[row]['a7_max'] != 'na':
            self.a7_max = int(float(df.loc[row]['a7_max']))
        else:
            self.a7_max = 'na'
        if self.a7_min == 'na' and self.a7_max == 'na':
            self.a7 = 'Evaluarea mărimii populației speciei se face pentru prima dată'
        else:
            self.a7 = str(self.a7_min) + ' - ' + str(self.a7_max) + self.unitate
        logging.debug('A setat A7'+' = ' +str(self.a7))
    #A.8	Mărimea populaţiei de referinţă pentru starea favorabilă în aria naturală protejată (df)
        if df.loc[row]['a8'] != 'na':
            self.a8 = int(float(df.loc[row]['a8']))
        else:
            self.a8 = 'Nu a putut fi stabilită o mărime de referință'
        logging.debug('A setat A8'+' = ' +str(self.a8))
    #A.9	Metodologia de apreciere a mărimii populaţiei de referinţă pentru starea favorabilă (df)
        self.a9 = df.loc[row]['a9']
        logging.debug('A setat A9'+' = ' +str(self.a9))
    #A.10	Raportul dintre mărimea populaţiei de referinţă pentru starea favorabilă şi mărimea populaţiei actuale (alg)
        if self.a8 == 'Nu a putut fi stabilită o mărime de referință' and self.pop_sit != 'na':
            self.a10 = rap[0]
        elif self.pop_sit == 0:
            self.a10 = rap[4]
        elif self.pop_sit == 'na':
            self.a10 = rap[0]
        else:
            r = self.a8/self.pop_sit
            if r<1:
                self.a10 = rap[1]
            elif r==1:
                self.a10 = rap[2]
            elif 1.25>r>1:
                self.a10 = rap[3]
            else:
                self.a10 = rap[4]

        logging.debug('A setat A10'+' = ' +str(self.a10))
    #A.11	Tendinţa actuală a mărimii populaţiei speciei (alg)   
   
        #if self.a7_min == 'na':
            #if self.pop_spa == 'na':
             #   self.a11 = ten[0]
            #elif self.pop_spa == self.pop_sit:
            #    self.a11 = ten[2]
            #elif self.pop_spa < self.pop_sit:
            #    self.a11 = ten[3]
            #else:
            #    self.a11 = ten[1]
        #elif int((self.a7_min+self.a7_max)/2) == self.pop_sit:
         #   self.a11 = ten[2]
        #elif int((self.a7_min+self.a7_max)/2) < self.pop_sit:
        #    self.a11 = ten[3]
        #else:
         #   self.a11 = ten[1]
        if df.loc[row]['a11'] == 'na':
            self.a11 = 'na'
        else:    
            self.a11 = ten[int(df.loc[row]['a11'])]
        logging.debug('A setat A11'+' = ' +str(self.a11))
    #A.12	Calitatea datelor privind tendinţa actuală a mărimii populaţiei speciei (cst) 
        if df.loc[row]['a12'] == 'na':
            self.a12 = 'na'
        else:
            self.a12 = cal[int(df.loc[row]['a12'])]
        logging.debug('A setat A12'+' = ' +str(self.a12))
    #A.13	Magnitudinea tendinţei actuale a mărimii populaţiei speciei  (cst) 
        self.a13_min = int(float(df.loc[row]['a13_min'])) if df.loc[row]['a13_min'] != 'na' else 'na'
        self.a13_max = int(float(df.loc[row]['a13_max'])) if df.loc[row]['a13_max'] != 'na' else 'na'
        if self.a13_min == 'na' and self.a13_max == 'na':
            self.a13 = 'Nu există suficiente informaţii pentru a putea aprecia magnitudinea tendinţei actuale a mărimii populaţiei speciei'
        else: 
            self.a13 = str(self.a13_min)+'%' + ' - ' + str(self.a13_max) + '%'
        logging.debug('A setat A13'+' = ' +str(self.a13))
    #A.14	Magnitudinea tendinţei actuale a mărimii populaţiei speciei exprimată prin calificative    (cst)
        self.a14 = df.loc[row]['a14']
        logging.debug('A setat A14'+' = ' +str(self.a14))
    #A.15	Structura populaţiei speciei    (cst)
        self.a15 = 'Nu există date privind structura populaţiei.'
        logging.debug('A setat A15'+' = ' +str(self.a15))
    #A.16	Starea de conservare din punct de vedere al populaţiei speciei  (alg)  
        if self.a8 == 'Nu a putut fi stabilită o mărime de referință':
            self.a16 = con['X']
            self.col_cons_pop = 3
        elif self.a10 in [rap[1], rap[2]]:
            self.a16 = con['FV']
            self.col_cons_pop = 0
        elif self.a13 != 'Nu există suficiente informaţii pentru a putea aprecia magnitudinea tendinţei actuale a mărimii populaţiei speciei':
            if self.a14 == '>5%' and self.a10 in [rap[4], rap[3]]:
                self.a16 = con['U2']
                self.col_cons_pop = 2
            else:
                self.a16 = con['U1']
                self.col_cons_pop = 1                
        elif self.a14 != 'Nu există suficiente informaţii pentru a putea aprecia magnitudinea tendinţei actuale a mărimii populaţiei speciei':
            if self.a14 == '>5%' and self.a10 in [rap[4], rap[3]]:
                self.a16 = con['U2']
                self.col_cons_pop = 2 
            else:
                self.a16 = con['U1']
                self.col_cons_pop = 1               
        else:
            self.a16 = con['U1']
            self.col_cons_pop = 1
        self.txt_cons_pop = 'A10: '+ self.a10 + ', A15: '+self.a15
        logging.debug('A setat A16'+' = ' +str(self.a16))
    #A.17	Tendinţa stării de conservare din punct de vedere al populaţiei speciei (cst)
        self.a17 = ten_con[df.loc[row]['a17']] 
        logging.debug('A setat A17'+' = ' +str(self.a17))
    #A.18	Starea de conservare necunoscută din punct de vedere al populaţiei  (alg) 
        self.a18 = df.loc[row]['a18'] if self.a16 == con['X'] else 'A fost calculată starea de conservare din punctul de vedere al populației speciei'
        logging.debug('A setat A18'+' = ' +str(self.a18))    



            #Variabilele pentru tabelul de conservare B
    #B.3	Suprafaţa habitatului speciei în aria naturală protejată (df + alg)
        self.b3_min = int(float(df.loc[row]['b3_min'])) if df.loc[row]['b3_min'] != 'na' else 'na'
        self.b3_max = int(float(df.loc[row]['b3_max'])) if df.loc[row]['b3_max'] != 'na' else 'na'
        if self.b3_min != 'na' and self.b3_max != 'na':
            self.b3 = str(self.b3_min)+' - '+str(self.b3_max)+' ha'
            self.hab = int(float((self.b3_min+self.b3_max)/2)) #Media habitatului
        else:
            self.b3 = 'Suprafața habitatului nu este disponibilă'
            self.hab = 'na' #Media habitatului  
        logging.debug('A setat B3'+' = ' +str(self.hab))
    #B.4	Calitatea datelor pentru suprafaţa habitatului speciei (df)
        self.b4 = cal[df.loc[row]['b4']]
        logging.debug('A setat B4'+' = ' +str(self.b4))
    #B.5	Suprafaţa reevaluată a habitatului speciei din planul de management anterior (cst)
        self.b5_min = int(float(df.loc[row]['b5_min'])) if df.loc[row]['b5_min'] != 'na' else 'na'
        self.b5_max = int(float(df.loc[row]['b5_max'])) if df.loc[row]['b5_max'] != 'na' else 'na'
        if self.b5_min == 'na' and self.b5_max == 'na':
            self.b5 = 'Evaluarea suprafeţei habitatului speciei se face pentru prima dată'
        else:
            self.b5 =  str(self.b5_min) + ' - ' + str(self.b5_max)
        logging.debug('A setat B5'+' = ' +str(self.b5))
    #B.6	Suprafaţa  adecvată a habitatului speciei în aria naturală protejată (cst)
        if df.loc[row]['b6'] == 'na':
            self.b6 = 'Nu există date suficiente pentru evaluarea suprafeței adecvate a habitatului speciei'
        else:
            self.b6 = int(float(df.loc[row]['b6']))
        logging.debug('A setat B6'+' = ' +str(self.b6))
    #B.7	Metodologia de apreciere a suprafeţei  adecvate a habitatului speciei în aria naturală protejată (cst)
        if self.b6 == 'Nu există date suficiente pentru evaluarea suprafeței adecvate a habitatului speciei':
            self.b7 = 'Nu a fost estimată suprafața adecvată a habitatului speciei'
        else:
            self.b7 = df.loc[row]['b7']
        logging.debug('A setat B7'+ ' = ' + str(self.b7))
    #B.8	Raportul dintre suprafaţa adecvată a habitatului speciei şi suprafaţa actuală a habitatului speciei (df)
        if self.b6 == 'Nu există date suficiente pentru evaluarea suprafeței adecvate a habitatului speciei':
            self.b8 = rap[df.loc[row]['b8']]
            self.b88 = rap[df.loc[row]['b8']] #Pentru a fi folosit in matricea starii de conservare
        else:
            self.b8 = 'Nu este cazul'
            self.b88 = rap[df.loc[row]['b8']]
        logging.debug('A setat B8'+' = ' +str(self.b8))
    #B.9	Tendinţa actuală a suprafeţei habitatului speciei (cst)
        self.b9 = ten[df.loc[row]['b9']]
        logging.debug('A setat B9'+' = ' +str(self.b9))
    #B.10	Calitatea datelor privind tendinţa actuală a suprafeţei habitatului speciei (cst)
        self.b10 = cal[int(df.loc[row]['b10'])] if df.loc[row]['b10'] != 'na' else 'na'
        logging.debug('A setat B10'+' = ' +str(self.b10))
    #B.11	Calitatea habitatului speciei în aria naturală protejată (df)
        self.b11 = df.loc[row]['b11']
        logging.debug('A setat B11'+' = ' +str(self.b11))
    #B.12	Tendinţa actuală a calităţii habitatului speciei (cst)
        self.b12 = ten[df.loc[row]['b12']]
        logging.debug('A setat B12'+' = ' +str(self.b12))
    #B.13	Calitatea datelor privind tendinţa actuală a calităţii habitatului speciei (cst)
        self.b13 = cal[int(df.loc[row]['b13'])] if df.loc[row]['b13'] != 'na' else 'na'
        logging.debug('A setat B13'+ ' = ' +str(self.b13))
    #B.14	Tendinţa actuală globală a habitatului speciei funcţie de tendinţa suprafeţei şi de tendinţa calităţii habitatului speciei (cst)
        if self.b9 == ten[0] or self.b12 == ten[0]:
            self.b14 = ten[0]
        elif [self.b9, self.b12] == [ten[2],ten[2]] :
            self.b14 = ten[2]
        elif [self.b9, self.b12] == [ten[2],ten[3]] or [self.b9, self.b12] == [ten[3],ten[3]]:
            self.b14 = ten[3]
        elif [self.b9, self.b12] == [ten[2],ten[1]] or [self.b9, self.b12] == [ten[1],ten[1]]:
            self.b14 = ten[1]
        else:
            self.b14 = ten[0]
        logging.debug('A setat B14'+' = ' +str(self.b14))
    #B.15	Starea de conservare din punct de vedere al habitatului speciei (alg)
        if self.b6 == 'Nu există date suficiente pentru evaluarea suprafeței adecvate a habitatului speciei':
            if self.b88 == rap[4] or self.b11=='rea':
                self.b15 = con['U2']
                self.col_cons_hab = 2
            elif self.b88 in [rap[1],rap[2]] and self.b9 in [ten[3],ten[2]] and self.b11 in ['bună (adecvată)', 'medie']:
                self.b15 = con['FV']
                self.col_cons_hab = 0 
            elif self.b88 == rap[0] or self.hab == 'na':
                self.b15 = con['X']
                self.col_cons_hab = 3          
            else:
                self.b15 = con['U1']
                self.col_cons_hab = 1
        else:
            if (self.hab*100/self.b6) <= 75 or self.b11=='rea':
                self.b15 = con['U2']
                self.col_cons_hab = 2
            elif self.b3_max >= self.b6 >= self.b3_min and self.b9 in [ten[3],ten[2]] and self.b11 in ['bună (adecvată)', 'medie']:
                self.b15 = con['FV']
                self.col_cons_hab = 0                     
            else:
                self.b15 = con['U1']
                self.col_cons_hab = 1
                print('This is self.hab ', self.hab)

        
        #Second variant
#        if self.b88 in [rap[1],rap[2]] and self.b9 in [ten[3],ten[2]] and self.b11 in ['bună (adecvată)', 'medie']:
 #           self.b15 = con['FV']
  #          self.col_cons_hab = 0
   #     elif self.b88 == rap[4] or self.b11=='rea':
    #        self.b15 = con['U2']
     #
     #   elif self.b6 == 'Nu există date suficiente pentru evaluarea suprafeței adecvate a habitatului speciei':
       #     if (self.hab*100/self.b6) <= 75 or self.b11=='rea':
        #        self.b15 = con['U2']
         #       self.col_cons_hab = 2
        #else:
         #   self.b15 = con['U1']
          #  self.col_cons_hab = 1
        


        logging.debug('A setat B15'+' = ' +str(self.b15))
    #B.16	Tendinţa stării de conservare din punct de vedere al habitatului speciei (cst)
        if self.b15 in {con['U2'], con['U1']}:
            self.b16 = ten_con[df.loc[row]['b16']]
        else:
            self.b16 = 'Starea de conservare  din punct de vedere al habitatului speciei [B.15.] nu a fost evaluată ca nefavorabilă \
                - inadecvată sau nefavorabilă - rea'
        logging.debug('A setat B16'+' = ' +str(self.b16))
    #B.17	Starea de conservare necunoscută din punct de vedere al habitatului speciei (alg)
        if self.b15 == con['X']:
            self.b17 = df.loc[row]['b17']
        else:
            self.b17 = 'Nu este cazul'
        logging.debug('A setat B17'+' = ' +str(self.b17))


                                        
                                        #Variabilele pentru tabelul de conservare C
    #C.3	Tendinţa viitoare a mărimii populaţiei (df)
        self.c3 = ten[df.loc[row]['c3']]
        logging.debug('A setat C3'+' = ' +str(self.c3))
    #C.4	Raportul dintre mărimea populaţiei de referinţă pentru starea favorabilă şi mărimea populaţiei viitoare a speciei (df)
        self.c4 = rap[df.loc[row]['c4']]
        logging.debug('A setat C4'+' = ' +str(self.c4))
    #C.5	Perspectivele speciei din punct de vedere al populaţiei (alg)
        if self.a10 == rap[1] and self.c3 == ten[1] and self.c4 == rap[1]:
            self.c5 = persp['FV']
            self.mat4_14 = 'img_cons/fn2.png'
        elif self.a10 in [rap[3], rap[4]] and self.c3 == ten[3] and self.c4 == rap[1]:
            self.c5 = persp['FV']
            self.mat4_14 = 'img_cons/fn3.png'
        elif self.a10 in [rap[1], rap[2]] and self.c3 == ten[3] and self.c4 == rap[1]:
            self.c5 = persp['FV']
            self.mat4_14 = 'img_cons/fn4.png'
        elif self.a10 in [rap[1], rap[2]] and self.c3 == ten[2] and self.c4 in [rap[1], rap[2]]:
            self.c5 = persp['FV']
            self.mat4_14 = 'img_cons/fn4.png'
        elif self.a10 == rap[2] and self.c3 == ten[1] and self.c4 == rap[3]:
            self.c5 = persp['U1']
            self.mat4_14 = 'img_cons/fn1.png'
        elif self.a10 == rap[1] and self.c3 == ten[1] and self.c4 == rap[2]:
            self.c5 = persp['U1']
            self.mat4_14 = 'img_cons/fn2.png'
        elif self.a10 in [rap[3], rap[4]] and self.c3 == ten[3] and self.c4 == rap[2]:
            self.c5 = persp['U1']
            self.mat4_14 = 'img_cons/fn3.png'
        elif self.a10 in [rap[3], rap[4]] and self.c3 == ten[2] and self.c4 == rap[3]:
            self.c5 = persp['U1']
            self.mat4_14 = 'img_cons/fn4.png'
        elif self.a10 == rap[2] and self.c3 == ten[1] and self.c4 == rap[4]:
            self.c5 = persp['U2']
            self.mat4_14 = 'img_cons/fn1.png'
        elif self.a10 == rap[1] and self.c3 == ten[1] and self.c4 in [rap[4], rap[3]]:
            self.c5 = persp['U2']
            self.mat4_14 = 'img_cons/fn2.png'
        elif self.a10 in [rap[3], rap[4]] and self.c3 == ten[3] and self.c4 in [rap[4], rap[3]]:
            self.c5 = persp['U2']
            self.mat4_14 = 'img_cons/fn3.png'
        elif self.a10 in [rap[3], rap[4]] and self.c3 == ten[2] and self.c4 == rap[4]:
            self.c5 = persp['U2']
            self.mat4_14 = 'img_cons/fn4.png'
        else: #Here everything eles falls in unknown (the matrix doesnt cover all the cases)
            self.c5 = persp['X']
            self.mat4_14 = 'na'
        logging.debug('A setat C5'+' = ' +str(self.c5))
    #C.6	Tendinţa viitoare a suprafeţei habitatului speciei (df)    
        self.c6 = ten[df.loc[row]['c6']]
        logging.debug('A setat C6'+' = ' +str(self.c6))
    #C.7	Raportul dintre suprafaţa adecvată a habitatului speciei şi suprafaţa habitatului speciei în viitor (df)
        self.c7 = rap[df.loc[row]['c7']]
        logging.debug('A setat C7'+' = ' +str(self.c7))
    #C.8	Perspectivele speciei din punct de vedere al habitatului speciei (alg)
        if self.b88 == rap[1] and self.c6 == ten[1] and self.c7 == rap[1]:
            self.c8 = persp['FV']
        elif self.b88 in [rap[3], rap[4]] and self.c6 == ten[3] and self.c7 == rap[1]:
            self.c8 = persp['FV']
        elif self.b88 in [rap[1], rap[2]] and self.c6 == ten[3] and self.c7 == rap[1]:
            self.c8 = persp['FV']
        elif self.b88 in [rap[1], rap[2]] and self.c6 == ten[2] and self.c7 in [rap[1], rap[2]]:
            self.c8 = persp['FV']
        elif self.b88 == rap[2] and self.c6 == ten[1] and self.c7 == rap[3]:
            self.c8 = persp['U1']
        elif self.b88 == rap[1] and self.c6 == ten[1] and self.c7 == rap[2]:
            self.c8 = persp['U1']
        elif self.b88 in [rap[3], rap[4]] and self.c6 == ten[3] and self.c7 == rap[2]:
            self.c8 = persp['U1']
        elif self.b88 in [rap[3], rap[4]] and self.c6 == ten[2] and self.c7 == rap[3]:
            self.c8 = persp['U1']
        elif self.b88 == rap[2] and self.c6 == ten[1] and self.c7 == rap[4]:
            self.c8 = persp['U2']
        elif self.b88 == rap[1] and self.c6 == ten[1] and self.c7 in [rap[4], rap[3]]:
            self.c8 = persp['U2']
        elif self.b88 in [rap[3], rap[4]] and self.c6 == ten[3] and self.c7 in [rap[4], rap[3]]:
            self.c8 = persp['U2']
        elif self.b88 in [rap[3], rap[4]] and self.c6 == ten[2] and self.c7 == rap[4]:
            self.c8 = persp['U2']
        else: #Here everything eles falls in unknown (the matrix doesnt cover all the cases)
            self.c8 = persp['X']
        logging.debug('A setat C8'+' = ' +str(self.c8))      
    #C.9	Perspectivele speciei în viitor (alg)        
        if self.c8 == persp['FV'] and self.c5 == persp['FV']:
            self.c9 = persp['FV']
            self.col_cons_pre = 0
            self.mat6_col = 0 #Variabilele pentru matricea 6
            self.txt_mat6_col  = 'Perspective favorabile'
        elif self.c8 == persp['X'] and self.c5 == persp['X']:
            self.c9 = persp['X']
            self.col_cons_pre = 3
            self.mat6_col = 3 #Variabilele pentru matricea 6
            self.txt_mat6_col  = 'Perspective necunoscute'
        elif self.c8 == persp['U2'] or self.c5 == persp['U2']:
            self.c9 = persp['U2']
            self.col_cons_pre = 2 
            self.mat6_col = 2 #Variabilele pentru matricea 6
            self.txt_mat6_col  = 'Perspective nefavorabile - rele'
        else: 
            self.c9 =persp['U1']
            self.col_cons_pre = 1
            self.mat6_col = 1 #Variabilele pentru matricea 6
            self.txt_mat6_col  = 'Perspective nefavorabile - inadecvate' 
        logging.debug('A setat C9'+' = ' +str(self.c9))
    #C.10	Efectul cumulat al impacturilor asupra speciei în viitor (df)
        self.c10 = imp[df.loc[row]['c10']]
        logging.debug('A setat C10'+' = ' +str(self.c10))
    #C.11  Intensitatea presiunilor actuale asupra speciei
        cond_timep = impf['time'] == 'prezenta'
        cond_specp1 = impf['id_spec'] == str(self.cod_sp)+str(self.feno)
        cond_specp2 = impf['id_spec'] == 'id'
        self.c11s = impf.loc[(cond_specp1 | cond_specp2)&cond_timep&(impf['intensitate'] == 'scăzută')]['impact'].unique().tolist()
        self.c11m = impf.loc[(cond_specp1 | cond_specp2)&cond_timep&(impf['intensitate'] == 'medie')]['impact'].unique().tolist()
        self.c11r = impf.loc[(cond_specp1 | cond_specp2)&cond_timep&(impf['intensitate'] == 'ridicată')]['impact'].unique().tolist()
        self.c11_sin = impf.loc[(cond_specp1 | cond_specp2)&cond_timep]['cod'].unique().tolist()
        logging.debug('A setat C11s'+' = ' +str(self.c11s))
        logging.debug('A setat C11m'+' = ' +str(self.c11m))
        logging.debug('A setat C11r'+' = ' +str(self.c11r))
    #C.12  Intensitatea ameninţărilor viitoare asupra speciei
        cond_timev = impf['time'] == 'viitoare'
        cond_specv1 = impf['id_spec'] == str(self.cod_sp)+str(self.feno)
        cond_specv2 = impf['id_spec'] == 'id' 
        self.c12s = impf.loc[(cond_specv1 | cond_specv2)&cond_timev&(impf['intensitate'] == 'scăzută')]['impact'].unique().tolist()
        self.c12m = impf.loc[(cond_specv1 | cond_specv2)&cond_timev&(impf['intensitate'] == 'medie')]['impact'].unique().tolist()
        self.c12r = impf.loc[(cond_specv1 | cond_specv2)&cond_timev&(impf['intensitate'] == 'ridicată')]['impact'].unique().tolist()
        logging.debug('A setat C12s'+' = ' +str(self.c12s))
        logging.debug('A setat C12m'+' = ' +str(self.c12m))
        logging.debug('A setat C12r'+' = ' +str(self.c12r))
    #C.13  Viabilitatea pe termen lung a speciei
        self.c13 = df.loc[row]['c13']
        logging.debug('A setat C13'+' = ' +str(self.c13))
    #C.14  Starea de conservare din punct de vedere al perspectivelor speciei în viitor
        if self.c10 == imp['r'] or self.c9 == persp['U2'] or self.c13 == 'viabilitatea pe termen lung a speciei nu este asigurată':
            self.c14 = con['U2']
            self.mat6_col = 2
            self.txt_mat6_col = 'C10: ' + self.c10 + ', C9: ' +  self.c9 + ', C13: ' + self.c13
        elif self.c10 == imp['s'] and (self.c9 == persp['FV'] or self.c13 == 'viabilitatea pe termen lung a speciei este asigurată'):
            self.c14 = con['FV']
            self.mat6_col = 0
            self.txt_mat6_col = 'C10: ' + self.c10 + ', C9: ' +  self.c9 + ', C13: ' + self.c13
        elif self.c10 == imp['n'] and self.c9 == persp['X']:
            self.c14 = con['X']
            self.mat6_col = 3
            self.txt_mat6_col = 'C10: ' + self.c10 + ', C9: ' +  self.c9 + ', C13: ' + self.c13
        elif self.c10 == imp['n'] and self.c13 == 'nu există suficiente informaţii pentru a aprecia gradul de asigurare al viabilităţii pe termen lung al speciei':
            self.c14 = con['X']
            self.mat6_col = 3
            self.txt_mat6_col = 'C10: ' + self.c10 + ', C9: ' +  self.c9 + ', C13: ' + self.c13
        elif self.c9 == persp['X'] and self.c13 == 'nu există suficiente informaţii pentru a aprecia gradul de asigurare al viabilităţii pe termen lung al speciei':
            self.c14 = con['X']
            self.mat6_col = 3
            self.txt_mat6_col = 'C10: ' + self.c10 + ', C9: ' +  self.c9 + ', C13: ' + self.c13
        else:
            self.c14 = con['U1']
            self.mat6_col = 1
            self.txt_mat6_col = 'C10: ' + self.c10 + ', C9: ' +  self.c9 + ', C13: ' + self.c13       
        logging.debug('A setat C14'+' = ' +str(self.c14))
    #C.15  Tendinţa stării de conservare din punct de vedere al perspectivelor speciei în viitor
        self.c15 = ten_con[df.loc[row]['c15']]
        logging.debug('A setat C15'+' = ' +str(self.c15))
    #C.16  Starea de conservare necunoscută din punct de vedere al perspectivelor speciei în viitor
        self.c16 = df.loc[row]['c16']
        logging.debug('A setat C16'+' = ' +str(self.c16))




                                    #Variabilele pentru tabelul de conservare D
    #D.3. Starea globală de conservare a speciei (alg)
        if self.a16 == con['U2'] or self.b15 == con['U2'] or self.c14 == con['U2']:
            self.d3 = con['U2']
            self.d3_sin = 'NR'
            #Variabilele pentru tabelul matricea 7
            self.col_cons_glob = 2
            self.txt_cons_glob = 'A16: '+self.a16 +', B15: , '+ self.b15 +', C14: '+ self.c14
        elif (self.a16 == con['X'] and self.b15 == ten['X']) or (self.a16 == con['X'] and self.c14 == con['X']) or (self.b15 == con['X'] and self.c14 == con['X']):
            self.d3 = con['X']
            self.d3_sin = 'N'
            #Variabilele pentru tabelul matricea 7
            self.col_cons_glob = 3
            self.txt_cons_glob = 'A16: '+self.a16 +', B15: , '+ self.b15 +'C14: '+ self.c14                    
        elif self.b15 == con['FV'] and self.c14 == con['FV'] and self.a16 == con['FV']:
            self.d3 = con['FV']
            self.d3_sin = 'F'
            #Variabilele pentru tabelul matricea 7
            self.col_cons_glob = 0
            self.txt_cons_glob = 'A16: '+self.a16 +', B15: , '+ self.b15 +'C14: '+ self.c14           
        elif self.b15 == con['FV'] and self.c14 == con['FV'] and self.a16 == con['X']:
            self.d3 = con['FV']
            self.d3_sin = 'F'
            #Variabilele pentru tabelul matricea 7
            self.col_cons_glob = 0
            self.txt_cons_glob = 'A16: '+self.a16 +', B15: , '+ self.b15 +'C14: '+ self.c14                         
        elif self.b15 == con['FV'] and self.c14 == con['X'] and self.a16 == con['FV']:
            self.d3 = con['FV']
            self.d3_sin = 'F'
            #Variabilele pentru tabelul matricea 7
            self.col_cons_glob = 0
            self.txt_cons_glob = 'A16: '+self.a16 +', B15: , '+ self.b15 +'C14: '+ self.c14           
        elif self.b15 == con['X'] and self.c14 == con['FV'] and self.a16 == con['FV']:
            self.d3 = con['FV']
            self.d3_sin = 'F'
            #Variabilele pentru tabelul matricea 7
            self.col_cons_glob = 0
            self.txt_cons_glob = 'A16: '+self.a16 +', B15: , '+ self.b15 +'C14: '+ self.c14            
        else: 
            self.d3 = con['U1']
            self.d3_sin = 'NI'
            #Variabilele pentru tabelul matricea 7
            self.col_cons_glob = 1
            self.txt_cons_glob = 'A16: '+self.a16 +', B15: '+ self.b15 +',  C14: '+ self.c14    
        logging.debug('A setat D3'+' = ' +str(self.d3))        
    #D.4. Tendinţa stării globale de conservare a speciei (alg + df)
        if self.d3 in {con['U1'], con['U2']}:
            self.d4 = df.loc[row]['d4']
        else: self.d4 = 'Starea globală de conservare a speciei [D.3.] nu a fost evaluată ca nefavorabilă - inadecvată sau nefavorabilă - rea'
        logging.debug('A setat D4'+' = ' +str(self.d4)) 
    #D.5. Starea globală de conservare necunoscută (alg + df)
        if self.d3 == con['X']:
            self.d5 = df.loc[row]['d5']
        else: self.d5 = 'Starea globală de conservare a speciei [D.3.] nu a fost evaluată ca ”X” necunoscută.'
        logging.debug('A setat D5'+' = ' +str(self.d5))  
    #D.6. Informaţii suplimentare (df)
        self.d6 = df.loc[row]['d6']
        logging.debug('A setat D6'+' = ' +str(self.d6)) 
        logging.info('A fost finalizata initializarea clasei specie pentru '+str(self.lat_sp))

                        #ATRIBUTES FOR SINTETIC MATRIX

        #Condition for specie &feno
        cond_spec1 = impf['id_spec'] == str(self.cod_sp)+str(self.feno)
        #Condition for specie_feno - (when specie_feno is selected in "id_spec" will be "id")
        cond_spec2 = impf['id_spec'] == 'id'
        #Conditions for pressure time
        cond_timev = impf['time'] == 'viitoare'
        cond_timep = impf['time'] == 'prezenta'

   
        #List of the presures on specie/feno, both present and future 
        self.sp_impacts = impf.loc[cond_spec1 | cond_spec2 ]['impact_pr'].unique().tolist()
        #List of all impacts ids - used for making measures list measures
        self.sp_impactsIds = impf.loc[cond_spec1 | cond_spec2 ]['id'].unique().tolist()
        #List of the present impacts codes
        self.sp_impPcodes = impf.loc[(cond_spec1 | cond_spec2)&cond_timep]['cod'].unique().tolist()
        #List of the future impacts codes
        self.sp_impVcodes = impf.loc[(cond_spec1 | cond_spec2)&cond_timev]['cod'].unique().tolist()
        
        #Setting the conditions for measures 
        cond_specm1 = msf['id_spec'] == str(self.cod_sp)+str(self.feno)
        cond_specm2 = msf['id_spec'] == 'id'
        cond_specm3 = msf['id_presiune'].isin(self.sp_impactsIds)
        #List of measures for specie/feno
        self.ms_list = msf.loc[cond_specm1 | cond_specm2 | cond_specm3]['masura'].unique().tolist()




class site():
    def __init__(self, the_site):
        self.the_site = the_site

        #Getting the table for the site from db
    def bf(self):
        conn=sqlite3.connect('infosites.db')
        tbl = pd.read_sql_query('SELECT * FROM '+self.the_site, conn)
        return tbl

        #Getting the master table from gdrive
    def master(self, xls):
        df = pd.read_excel(xls, sheet_name='master')
        df = df.drop(df.index[0])
        df.fillna('na', inplace=True)
        df = df[df.id != 'na']
        return df

        #Getting the impacts table from gdrive
    def impacts(self, xls):
        df = pd.read_excel(xls, sheet_name='presiuni_sit')
        df.fillna('na', inplace=True)
        df = df[df.impact != 'na']
        return df

        #Getting the measures table from gdrive
    def masuri(self, xls):
        df = pd.read_excel(xls, sheet_name='masuri_sit')
        df.fillna('na', inplace=True)
        df = df[df.masura != 'na']
        return df

        #Getting the descriptions table from gdrive
    def descrieri(self, xls):
        df = pd.read_excel(xls, sheet_name='descrieri')
        df.fillna('na', inplace=True)
        df = df[df.den_latin != 'na']
        return df




